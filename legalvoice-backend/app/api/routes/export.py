import io
from typing import Any
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel
from app.core.security import get_current_user_id

router = APIRouter()


def _apply_marks(run, marks: list[dict]) -> None:
    for mark in marks:
        t = mark.get("type")
        if t == "bold":
            run.bold = True
        elif t == "italic":
            run.italic = True
        elif t == "underline":
            run.underline = True
        elif t == "strike":
            run.font.strike = True


def _add_paragraph(doc: DocxDocument, node: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for child in node.get("content", []):
        if child.get("type") == "text":
            run = p.add_run(child.get("text", ""))
            run.font.size = Pt(11)
            _apply_marks(run, child.get("marks", []))


def _add_heading(doc: DocxDocument, node: dict) -> None:
    level = node.get("attrs", {}).get("level", 1)
    style = f"Heading {min(level, 3)}"
    text = "".join(
        c.get("text", "") for c in node.get("content", []) if c.get("type") == "text"
    )
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x3D, 0x58)


def _add_list_item(doc: DocxDocument, node: dict, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    for child in node.get("content", []):
        if child.get("type") == "paragraph":
            for inline in child.get("content", []):
                if inline.get("type") == "text":
                    run = p.add_run(inline.get("text", ""))
                    run.font.size = Pt(11)
                    _apply_marks(run, inline.get("marks", []))


def _process_node(doc: DocxDocument, node: dict, ordered: bool = False, depth: int = 0) -> None:
    if depth > 50:
        return
    t = node.get("type")
    if t == "heading":
        _add_heading(doc, node)
    elif t == "paragraph":
        _add_paragraph(doc, node)
    elif t == "bulletList":
        for child in node.get("content", []):
            _add_list_item(doc, child, ordered=False)
    elif t == "orderedList":
        for child in node.get("content", []):
            _add_list_item(doc, child, ordered=True)
    elif t == "blockquote":
        for child in node.get("content", []):
            _process_node(doc, child, depth=depth + 1)
    elif t == "horizontalRule":
        doc.add_paragraph("─" * 60)


def content_to_docx(title: str, content: dict[str, Any]) -> bytes:
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(90)

    # Title
    title_para = doc.add_heading(title, level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x3D, 0x58)
        run.font.size = Pt(18)
    doc.add_paragraph()

    for node in content.get("content", []):
        _process_node(doc, node)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class ExportRequest(BaseModel):
    title: str
    content: dict[str, Any]


@router.post("/docx")
async def export_docx(
    body: ExportRequest,
    user_id: str = Depends(get_current_user_id),
):
    if not body.content.get("content"):
        raise HTTPException(status_code=400, detail="El documento está vacío")

    try:
        docx_bytes = content_to_docx(body.title, body.content)
    except Exception:
        raise HTTPException(status_code=500, detail="Error interno al generar el documento. Inténtalo de nuevo.")

    safe_title = "".join(c if c.isalnum() or c in " .-_" else "_" for c in body.title)[:80]
    filename = f"{safe_title}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
